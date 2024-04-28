#Librerias
import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter import font
from tkinter import messagebox
import docx
from docx import Document
import os

#Creacion de documento
doc = docx.Document()

#Creacion ventana
root = tk.Tk()

#Constantes
F = font.Font(size=16)

#Creacion ventana root y frame
frm = ttk.Frame(root, padding=10)
root.title("Etiquetas")
root.geometry("450x300")
root.resizable(False, False)
frm.pack()

def Salir():
    salir = messagebox.askquestion("Salir", "Desea salir?")
    if salir == "yes":
        root.destroy()
    else:
        return

#Funcion para agregar datos
def Agregar():
    if entrada_nombres.get() == "" or entrada_expediente.get() == "" or entrada_municipio.get() == "":
        messagebox.showerror("Error", "Falta un dato")
        return
    else:
        nombres.append(entrada_nombres.get())
        expediente.append(entrada_expediente.get())
        municipio.append(entrada_municipio.get())
        print(nombres)
        print(expediente)
        print(municipio)
        entrada_nombres.delete(0, END)
        entrada_expediente.delete(0, END)
        entrada_municipio.delete(0, END)

#Funcion para guardar en documento
def Guardar_documento():

    nombre_archivo = "documento.docx"
    incremental = 0
    idocumento = 0

    while os.path.exists(nombre_archivo):
        nombre_archivo = "documento" + str(idocumento) + ".docx"
        idocumento += 1

    while incremental < len(nombres):
        doc.add_paragraph(str(nombres[incremental]) + " " + str(expediente[incremental]))
        doc.add_paragraph(str(municipio[incremental]))
        doc.save(nombre_archivo)
        incremental += 1

#Etiqueta para nombres
NOMBRES = ttk.Label(frm, text="    Nombres y Apellidos    ", font=F).pack()

#Entry para Nombres
entrada_nombres = ttk.Entry(frm, width=30, justify=CENTER)
entrada_nombres.pack()

#Arreglo para guardar los nombres
nombres = []

#Etiqueta para numero de expediente
NO_EXPEDIENTE = ttk.Label(frm, text="No. de Expediente", font=F).pack()

#Entry para numero de expediente
entrada_expediente = ttk.Entry(frm, width=10, justify=CENTER)
entrada_expediente.pack()

#Arreglo para guardar el numero de expediente
expediente = []

#Etiqueta para municipio
MUNICIPIO = ttk.Label(frm, text="Municipio", font=F).pack()

#Entry para municipio
entrada_municipio = ttk.Entry(frm, width=30, justify=CENTER)
entrada_municipio.pack()

#Arreglo para guardar el municipio
municipio = []

#Boton para guardar en arreglos
BOTON_AGREGAR = ttk.Button(frm, text="Guardar", command=Agregar).pack()

#Boton para guardar en documento
BOTON_DOCUMENTO = ttk.Button(frm, text="Guardar en documento", command=Guardar_documento).pack()

#Boton para salir
BOTON_SALIR = ttk.Button(frm, text="Salir", command=Salir).pack()

root.mainloop()
